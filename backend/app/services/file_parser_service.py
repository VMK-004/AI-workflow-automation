"""File parser service for extracting text from various document formats"""
import logging
from typing import List, Dict, Any, Optional
from io import BytesIO
import asyncio

logger = logging.getLogger(__name__)


class FileParserService:
    """Service for parsing documents and extracting text"""
    
    # Maximum file size (10 MB)
    MAX_FILE_SIZE = 10 * 1024 * 1024
    
    @staticmethod
    async def parse_file(
        filename: str,
        content: bytes,
        chunk_size: int = 1000,
        chunk_overlap: int = 200
    ) -> List[Dict[str, Any]]:
        """
        Parse a file and extract text, optionally chunking it.
        
        Args:
            filename: Name of the file
            content: File content as bytes
            chunk_size: Size of text chunks (in characters, approximate)
            chunk_overlap: Overlap between chunks (in characters)
        
        Returns:
            List of document dicts with 'text' and 'metadata' keys
        """
        # Check file size
        if len(content) > FileParserService.MAX_FILE_SIZE:
            raise ValueError(f"File size ({len(content)} bytes) exceeds maximum ({FileParserService.MAX_FILE_SIZE} bytes)")
        
        # Determine file type from extension
        file_ext = filename.lower().split('.')[-1] if '.' in filename else ''
        
        # Extract text based on file type
        if file_ext == 'pdf':
            text = await FileParserService._parse_pdf(content)
        elif file_ext in ['docx', 'doc']:
            text = await FileParserService._parse_docx(content)
        elif file_ext in ['txt', 'md', 'text']:
            text = await FileParserService._parse_text(content)
        else:
            raise ValueError(f"Unsupported file type: {file_ext}. Supported: pdf, docx, txt, md")
        
        if not text or not text.strip():
            raise ValueError("No text content extracted from file")
        
        # Chunk the text if it's long enough
        chunks = FileParserService._chunk_text(text, chunk_size, chunk_overlap)
        
        # Create document dicts with metadata
        documents = []
        for i, chunk in enumerate(chunks):
            documents.append({
                "text": chunk.strip(),
                "metadata": {
                    "source_file": filename,
                    "chunk_index": i,
                    "total_chunks": len(chunks)
                }
            })
        
        logger.info(f"Parsed {filename}: extracted {len(chunks)} chunks from {len(text)} characters")
        
        return documents
    
    @staticmethod
    async def _parse_pdf(content: bytes) -> str:
        """Extract text from PDF file"""
        try:
            import PyPDF2
        except ImportError:
            raise ImportError("PyPDF2 is not installed. Install it with: pip install PyPDF2")
        
        # Run in executor to avoid blocking
        loop = asyncio.get_event_loop()
        text = await loop.run_in_executor(
            None,
            FileParserService._extract_pdf_text,
            content
        )
        return text
    
    @staticmethod
    def _extract_pdf_text(content: bytes) -> str:
        """Synchronous PDF text extraction"""
        import PyPDF2
        
        text_parts = []
        pdf_file = BytesIO(content)
        
        try:
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                except Exception as e:
                    logger.warning(f"Failed to extract text from PDF page {page_num}: {str(e)}")
            
            return "\n\n".join(text_parts)
        except Exception as e:
            logger.error(f"Failed to parse PDF: {str(e)}")
            raise ValueError(f"Failed to parse PDF file: {str(e)}")
    
    @staticmethod
    async def _parse_docx(content: bytes) -> str:
        """Extract text from DOCX file"""
        try:
            from docx import Document
        except ImportError:
            raise ImportError("python-docx is not installed. Install it with: pip install python-docx")
        
        # Run in executor to avoid blocking
        loop = asyncio.get_event_loop()
        text = await loop.run_in_executor(
            None,
            FileParserService._extract_docx_text,
            content
        )
        return text
    
    @staticmethod
    def _extract_docx_text(content: bytes) -> str:
        """Synchronous DOCX text extraction"""
        from docx import Document
        
        try:
            doc_file = BytesIO(content)
            doc = Document(doc_file)
            
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                    if row_text:
                        text_parts.append(row_text)
            
            return "\n\n".join(text_parts)
        except Exception as e:
            logger.error(f"Failed to parse DOCX: {str(e)}")
            raise ValueError(f"Failed to parse DOCX file: {str(e)}")
    
    @staticmethod
    async def _parse_text(content: bytes) -> str:
        """Extract text from plain text file"""
        try:
            # Try UTF-8 first
            text = content.decode('utf-8')
        except UnicodeDecodeError:
            try:
                # Fallback to latin-1
                text = content.decode('latin-1')
            except UnicodeDecodeError:
                # Last resort: ignore errors
                text = content.decode('utf-8', errors='ignore')
        
        return text
    
    @staticmethod
    def _chunk_text(text: str, chunk_size: int = 1000, chunk_overlap: int = 200) -> List[str]:
        """
        Split text into overlapping chunks.
        
        Args:
            text: Text to chunk
            chunk_size: Target chunk size in characters
            chunk_overlap: Overlap between chunks in characters
        
        Returns:
            List of text chunks
        """
        if len(text) <= chunk_size:
            return [text]
        
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + chunk_size
            
            # Try to break at sentence boundary if possible
            if end < len(text):
                # Look for sentence endings
                last_period = text.rfind('.', start, end)
                last_newline = text.rfind('\n', start, end)
                
                # Prefer sentence boundary, then paragraph boundary
                break_point = max(last_period, last_newline)
                
                if break_point > start + chunk_size // 2:  # Only break if we have enough content
                    end = break_point + 1
            
            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)
            
            # Move start forward with overlap
            start = end - chunk_overlap
            if start >= len(text):
                break
        
        return chunks





